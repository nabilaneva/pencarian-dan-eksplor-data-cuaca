from flask import Flask, render_template, request, send_file, jsonify, flash, session, redirect
import os
import requests
import pythoncom
from docx import Document
from docx2pdf import convert

app = Flask(__name__)
app.secret_key = 'kunci_rahasia_anda'

# Fungsi untuk mendapatkan data dari API BMKG
def get_data_from_bmkg():
    url = "https://stamet-juanda.bmkg.go.id/cuwis/json/test.json"
    headers = {'User-Agent': 'AplikasiCuaca/1.0'}
    try:
        response = requests.get(url, headers=headers)
        response.raise_for_status()
        return response.json()  # Mengembalikan data JSON
    except requests.exceptions.RequestException as e:
        flash(f"Gagal terhubung ke API BMKG. Kesalahan: {str(e)}")
        return None

# Route untuk halaman utama
@app.route('/', methods=['GET', 'POST'])
def index():
    bmkg_data = get_data_from_bmkg()
    data_sudah_ditemukan = False  # Default ke False, artinya belum ada pencarian yang berhasil

    # Dapatkan daftar unik kota dari API, lalu urutkan berdasarkan abjad
    daftar_kota = []
    if bmkg_data:
        daftar_kota = sorted(list(set(item['Kota'] for item in bmkg_data)))  # Menghapus duplikasi kota dan mengurutkan

    if 'data_cuaca_terdahulu' not in session:
        session['data_cuaca_terdahulu'] = []  # Inisialisasi session untuk data cuaca

    # Mengecek apakah jumlah pencarian sudah mencapai batas
    if len(session['data_cuaca_terdahulu']) >= 5:
        flash("Batas pencarian maksimum tercapai. Hapus beberapa hasil pencarian untuk melanjutkan.")
        return render_template('index.html', daftar_kota=daftar_kota, data_cuaca_terdahulu=session['data_cuaca_terdahulu'], data_sudah_ditemukan=data_sudah_ditemukan)

    if request.method == 'POST':
        tanggal = request.form['tanggal']
        kota = request.form['kota']
        kecamatan = request.form['kecamatan']
        data_cuaca = None

        if bmkg_data:
            # Temukan data yang sesuai berdasarkan input
            for item in bmkg_data:
                if item['Tanggal'] == tanggal and item['Kota'] == kota and item['Kecamatan'] == kecamatan:
                    data_cuaca = item
                    data_sudah_ditemukan = True
                    break

            if data_cuaca:
                # Tambahkan data cuaca baru ke session tanpa menghapus yang sebelumnya
                session['data_cuaca_terdahulu'].append(data_cuaca)
                session.modified = True  # Untuk memastikan session diperbarui
            else:
                flash("Data tidak ditemukan untuk kota, kecamatan, atau tanggal yang dipilih.")
        else:
            flash("Gagal mendapatkan data dari API BMKG.")

    return render_template('index.html', daftar_kota=daftar_kota, data_cuaca_terdahulu=session['data_cuaca_terdahulu'], data_sudah_ditemukan=data_sudah_ditemukan)

# Route untuk mendapatkan kecamatan berdasarkan kota
@app.route('/get_kecamatan/<kota>', methods=['GET'])
def get_kecamatan(kota):
    bmkg_data = get_data_from_bmkg()
    if bmkg_data:
        kecamatan_set = set(item['Kecamatan'] for item in bmkg_data if item['Kota'] == kota)
        kecamatan_sorted = sorted(list(kecamatan_set))  # Urutkan kecamatan berdasarkan abjad
        return jsonify({"kecamatan": kecamatan_sorted})
    else:
        return jsonify({"kecamatan": []})

# Fungsi untuk mengganti placeholder dalam dokumen Word
def replace_text(doc_obj, placeholder, replacement):
    for paragraph in doc_obj.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, replacement)
    
    # Periksa juga tabel (jika template memiliki tabel dengan placeholder)
    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                if placeholder in cell.text:
                    cell.text = cell.text.replace(placeholder, replacement)

# Route untuk menghapus pencarian berdasarkan ID atau kriteria tertentu
@app.route('/delete_search', methods=['POST'])
def delete_search():
    search_id = request.json.get('search_id')  # ID dari pencarian yang akan dihapus

    # Pastikan 'data_cuaca_terdahulu' ada di session
    if 'data_cuaca_terdahulu' in session:
        # Hapus data yang sesuai dengan ID
        session['data_cuaca_terdahulu'] = [
            data for data in session['data_cuaca_terdahulu']
            if data.get('ID') != search_id
        ]
        session.modified = True
        return jsonify({"success": True})
    return jsonify({"success": False, "message": "Data tidak ditemukan."})

# Route untuk ekspor data cuaca ke Word
@app.route('/export', methods=['POST'])
def export_to_word():
    # Ambil data dari form
    data_cuaca = request.form

    # Buka dokumen template
    template_path = os.path.join(os.path.dirname(__file__), 'templates', 'draft laporan.docx')  # Sesuaikan dengan path template yang diunggah
    doc = Document(template_path)

    # Mengganti placeholder di template sesuai data cuaca
    replace_text(doc, '{{tanggal}}', data_cuaca.get('Tanggal', 'Tanggal Tidak Tersedia'))
    replace_text(doc, '{{kota}}', data_cuaca.get('Kota', 'Kota Tidak Tersedia'))
    replace_text(doc, '{{kecamatan}}', data_cuaca.get('Kecamatan', 'Kecamatan Tidak Tersedia'))
    replace_text(doc, '{{suhu_min}}', data_cuaca.get('Suhu_Min', 'Suhu Tidak Tersedia'))
    replace_text(doc, '{{suhu_max}}', data_cuaca.get('Suhu_Max', 'Suhu Tidak Tersedia'))
    replace_text(doc, '{{kelembaban_min}}', data_cuaca.get('Kelembapan_Min', 'Kelembaban Tidak Tersedia'))
    replace_text(doc, '{{kelembaban_max}}', data_cuaca.get('Kelembapan_Max', 'Kelembaban Tidak Tersedia'))
    replace_text(doc, '{{arah_angin}}', data_cuaca.get('Arah_Angin', 'Angin Tidak Tersedia'))
    replace_text(doc, '{{cuaca_pagi}}', data_cuaca.get('Cuaca_Pagi', 'Data Tidak Tersedia'))
    replace_text(doc, '{{cuaca_siang}}', data_cuaca.get('Cuaca_Siang', 'Data Tidak Tersedia'))
    replace_text(doc, '{{cuaca_malam}}', data_cuaca.get('Cuaca_Malam', 'Data Tidak Tersedia'))
    replace_text(doc, '{{cuaca_dini}}', data_cuaca.get('Cuaca_Dini', 'Data Tidak Tersedia'))

    # Simpan dokumen yang sudah diisi
    os.makedirs("output", exist_ok=True)
    output_word_path = os.path.join("output", "laporan_cuaca.docx")
    doc.save(output_word_path)

    # Kirim file Word ke pengguna
    return send_file(output_word_path, as_attachment=True)

# Route untuk ekspor data cuaca ke PDF (mengonversi dari Word)
@app.route('/export/pdf', methods=['POST'])
def export_to_pdf():
    # Jalankan fungsi yang sama seperti export_to_word untuk menghasilkan dokumen Word terlebih dahulu
    data_cuaca = request.form
    output_word_path = os.path.join("output", "laporan_cuaca.docx")

    # Pastikan file Word sudah dibuat
    if not os.path.exists(output_word_path):
        export_to_word()  # Panggil fungsi jika file belum ada

    # Tentukan path PDF
    output_pdf_path = os.path.join("output", "laporan_cuaca.pdf")

    #Inisialisasi COM sebelum konversi dan uninitialize setelahnya
    pythoncom.CoInitialize()
    try:
        # Konversi Word ke PDF
        convert(output_word_path, output_pdf_path)
    finally:
        pythoncom.CoUninitialize()

    # Kirim file PDF ke pengguna
    return send_file(output_pdf_path, as_attachment=True)

@app.route('/delete_all_searches', methods=['POST'])
def delete_all_searches():
    if 'data_cuaca_terdahulu' in session:
        session['data_cuaca_terdahulu'] = []  # Kosongkan seluruh data pencarian
        session.modified = True  # Pastikan session diperbarui
        return jsonify({"success": True})
    return jsonify({"success": False, "message": "Data tidak ditemukan."})

@app.route('/export_all_word', methods=['GET'])
def export_all_word():
    if 'data_cuaca_terdahulu' not in session or not session['data_cuaca_terdahulu']:
        flash("Tidak ada data untuk diekspor.")
        return redirect('/')
    
    # Menentukan template berdasarkan jumlah data
    jumlah_data = len(session['data_cuaca_terdahulu'])
    template_path = os.path.join(os.path.dirname(__file__), 'templates', f'draft {jumlah_data}.docx')

    # Buat dokumen Word berdasarkan template
    doc = Document(template_path)
    for i, data in enumerate(session['data_cuaca_terdahulu'], 1):
        replace_text(doc, f'{{{{tanggal_{i}}}}}', str(data.get('Tanggal', 'Tidak Tersedia')))
        replace_text(doc, f'{{{{kota_{i}}}}}', str(data.get('Kota', 'Tidak Tersedia')))
        replace_text(doc, f'{{{{kecamatan_{i}}}}}', str(data.get('Kecamatan', 'Tidak Tersedia')))
        replace_text(doc, f'{{{{suhu_min_{i}}}}}', str(data.get('Suhu_Min', 'Tidak Tersedia')))
        replace_text(doc, f'{{{{suhu_max_{i}}}}}', str(data.get('Suhu_Max', 'Tidak Tersedia')))
        replace_text(doc, f'{{{{kelembaban_min_{i}}}}}', str(data.get('Kelembapan_Min', 'Tidak Tersedia')))
        replace_text(doc, f'{{{{kelembaban_max_{i}}}}}', str(data.get('Kelembapan_Max', 'Tidak Tersedia')))
        replace_text(doc, f'{{{{arah_angin_{i}}}}}', str(data.get('Arah_Angin', 'Tidak Tersedia')))
        replace_text(doc, f'{{{{cuaca_pagi_{i}}}}}', str(data.get('Cuaca_Pagi', 'Tidak Tersedia')))
        replace_text(doc, f'{{{{cuaca_siang_{i}}}}}', str(data.get('Cuaca_Siang', 'Tidak Tersedia')))
        replace_text(doc, f'{{{{cuaca_malam_{i}}}}}', str(data.get('Cuaca_Malam', 'Tidak Tersedia')))
        replace_text(doc, f'{{{{cuaca_dini_{i}}}}}', str(data.get('Cuaca_Dini', 'Tidak Tersedia')))
        # Tambahkan replace_text untuk field lain sesuai kebutuhan

    # Simpan dokumen
    output_word_path = os.path.join("output", "laporan_cuaca_semua.docx")
    doc.save(output_word_path)

    # Kirim file ke pengguna
    return send_file(output_word_path, as_attachment=True)

@app.route('/export_all_pdf', methods=['GET'])
def export_all_pdf():
    # Pastikan file Word sudah dibuat
    export_all_word()  # Buat dokumen Word terlebih dahulu

    # Konversi dokumen Word ke PDF
    output_word_path = os.path.join("output", "laporan_cuaca_semua.docx")
    output_pdf_path = os.path.join("output", "laporan_cuaca_semua.pdf")
    
    # Inisialisasi COM sebelum konversi dan uninitialize setelahnya
    pythoncom.CoInitialize()
    try:
        convert(output_word_path, output_pdf_path)
    finally:
        pythoncom.CoUninitialize()

    # Kirim file PDF ke pengguna
    return send_file(output_pdf_path, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)